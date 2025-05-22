import os
import logging
import mimetypes
from pathlib import Path

# Specific file format libraries
import fitz  # PyMuPDF for PDF files
import docx  # python-docx for DOCX files

# Configure logging
logger = logging.getLogger('resume_analyzer.utils')

def detect_file_type(file_path):
    """
    Detect file type using file extension and MIME type.
    Returns a tuple of (extension, mime_type)
    """
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    mime_type, _ = mimetypes.guess_type(file_path)
    
    logger.debug(f"File {file_path}: extension={file_ext}, mime_type={mime_type}")
    
    return file_ext, mime_type

def validate_extracted_text(text):
    """
    Validate the extracted text to ensure it's not empty or corrupted.
    Returns cleaned text.
    """
    if not text:
        return ""
    
    # Remove extremely long words (likely parsing errors)
    cleaned_text = ' '.join(word for word in text.split() if len(word) < 50)
    
    # Check if text is too short
    if len(cleaned_text.strip()) < 10:
        logger.warning(f"Extracted text is very short ({len(cleaned_text)} chars)")
    
    return cleaned_text

def extract_text_from_pdf(file_path):
    """Extract text from PDF files using PyMuPDF"""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX files using python-docx"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")

def extract_text_from_doc(file_path):
    """Handle DOC files by suggesting conversion to DOCX"""
    logger.warning(f"DOC format detected for {file_path}. Direct DOC support is not available.")
    message = (
        "Old DOC format detected. For best results, please convert your file to DOCX format "
        "using Microsoft Word or LibreOffice and upload again."
    )
    raise ValueError(message)

def extract_text_from_txt(file_path):
    """Extract text from plain text files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            return file.read()
    except Exception as e:
        try:
            # Try alternate encoding if utf-8 fails
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e2:
            logger.error(f"Error extracting text from TXT {file_path}: {e2}")
            raise ValueError(f"Could not extract text from TXT: {str(e2)}")

def extract_text_from_file(file_path):
    """
    Extract text from various file formats.
    Supports PDF, DOCX, DOC, and TXT files.
    
    Args:
        file_path: Path to the file
    
    Returns:
        Extracted text as string
    
    Raises:
        ValueError: If file format is unsupported or text extraction fails
    """
    if not os.path.exists(file_path):
        logger.error(f"File does not exist: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Detect file type
    file_ext, mime_type = detect_file_type(file_path)
    
    # Extract text based on file type
    try:
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            text = extract_text_from_docx(file_path)
        elif file_ext == 'doc':
            # Will raise a user-friendly error suggesting conversion
            text = extract_text_from_doc(file_path)
        elif file_ext == 'txt':
            text = extract_text_from_txt(file_path)
        else:
            logger.error(f"Unsupported file format: {file_ext}")
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Validate and clean the extracted text
        text = validate_extracted_text(text)
        
        # Log text length
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        raise ValueError(f"Error extracting text: {str(e)}")
